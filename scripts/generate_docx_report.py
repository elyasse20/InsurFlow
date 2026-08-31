#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to convert docs/rapport_pfa_insurflow.md into a professionally formatted
Microsoft Word document (docs/Rapport_PFA_InsurFlow_EMSI.docx) respecting all
official EMSI guidelines.
"""

import os
import re
import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Set inner margins (padding) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Set background color of a cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_borders(cell, top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC", sz="4"):
    """Set borders for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{top}"/>
        <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{left}"/>
        <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{bottom}"/>
        <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{right}"/>
    </w:tcBorders>
    '''
    tcPr.append(parse_xml(borders_xml))

def set_row_cant_split(row):
    """Prevent a row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_row_header(row):
    """Mark a row as repeated header across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def add_field(p, field_name):
    """Add a dynamic Word field (like PAGE or NUMPAGES)."""
    run = p.add_run()
    r = run._r
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> %s </w:instrText>' % (nsdecls('w'), field_name))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def style_paragraph(p, font_name="Times New Roman", font_size=12, bold=False, italic=False,
                    color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5, space_before=6, space_after=6, first_line_indent=0.8):
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    else:
        p.paragraph_format.first_line_indent = Cm(0)

def add_formatted_text(p, text, default_bold=False, default_italic=False,
                       font_name="Times New Roman", font_size=12, default_color=RGBColor(0, 0, 0)):
    """Parse inline markdown formatting (bold, italic, code, sub/sup, html linebreaks) and add runs."""
    # First handle <br> or <br/> tags
    lines = re.split(r'<br\s*/?>', text)
    for l_idx, sub_text in enumerate(lines):
        if l_idx > 0:
            # Add line break in Word
            p.add_run().add_break()
        
        tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$|<ins>[^<]+</ins>)', sub_text)
        for token in tokens:
            if not token:
                continue
            run = p.add_run()
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = default_color
            run.bold = default_bold
            run.italic = default_italic

            if token.startswith('**') and token.endswith('**'):
                run.text = token[2:-2]
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run.text = token[1:-1]
                run.italic = True
            elif token.startswith('`') and token.endswith('`'):
                run.text = token[1:-1]
                run.font.name = "Consolas"
                run.font.size = Pt(max(7.5, font_size - 1.0))
                run.font.color.rgb = RGBColor(180, 20, 20)
            elif token.startswith('$') and token.endswith('$'):
                run.text = token[1:-1]
                run.font.name = "Cambria Math"
                run.italic = True
            elif token.startswith('<ins>') and token.endswith('</ins>'):
                run.text = token[5:-6]
                run.underline = True
                run.bold = True
            else:
                clean = re.sub(r'<[^>]+>', '', token)
                run.text = clean

def build_docx_report(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = docx.Document()

    # 1. Page Setup & Margins: 2.5 cm everywhere (EMSI Official Standard)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.different_first_page_header_footer = True

        # Header (En-tête)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(4)
        hrun = hp.add_run("InsurFlow — Rapport de Projet de Fin d'Année (EMSI 4IIR)")
        hrun.font.name = "Times New Roman"
        hrun.font.size = Pt(9)
        hrun.font.italic = True
        hrun.font.color.rgb = RGBColor(100, 116, 139)

        # Footer (Pied de page)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fp.paragraph_format.space_before = Pt(4)
        frun1 = fp.add_run("École Marocaine des Sciences de l'Ingénieur (EMSI)")
        frun1.font.name = "Times New Roman"
        frun1.font.size = Pt(9)
        frun1.font.color.rgb = RGBColor(100, 116, 139)
        
        frun_tab = fp.add_run("\t\tPage ")
        frun_tab.font.name = "Times New Roman"
        frun_tab.font.size = Pt(9)
        frun_tab.font.color.rgb = RGBColor(100, 116, 139)
        add_field(fp, "PAGE")
        frun_of = fp.add_run(" / ")
        frun_of.font.name = "Times New Roman"
        frun_of.font.size = Pt(9)
        frun_of.font.color.rgb = RGBColor(100, 116, 139)
        add_field(fp, "NUMPAGES")

    # Set normal style font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(15, 23, 42)

    # 2. Build Cover Page (Page de Garde officielle EMSI)
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_top.paragraph_format.space_before = Pt(0)
    p_top.paragraph_format.space_after = Pt(2)
    r_inst = p_top.add_run("ÉCOLE MAROCAINE DES SCIENCES DE L'INGÉNIEUR\nEMSI — HONORIS UNITED UNIVERSITIES")
    r_inst.font.name = "Times New Roman"
    r_inst.font.size = Pt(13)
    r_inst.bold = True
    r_inst.font.color.rgb = RGBColor(180, 0, 0) # EMSI Red

    p_filiere = doc.add_paragraph()
    p_filiere.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_filiere.paragraph_format.space_before = Pt(2)
    p_filiere.paragraph_format.space_after = Pt(18)
    r_fil = p_filiere.add_run("Filière : Ingénierie Informatique et Réseaux (4IIR)")
    r_fil.font.name = "Times New Roman"
    r_fil.font.size = Pt(12)
    r_fil.font.italic = True
    r_fil.font.color.rgb = RGBColor(71, 85, 105)

    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.space_before = Pt(15)
    p_banner.paragraph_format.space_after = Pt(15)
    r_ban = p_banner.add_run("RAPPORT DE PROJET DE FIN D'ANNÉE")
    r_ban.font.name = "Times New Roman"
    r_ban.font.size = Pt(18)
    r_ban.bold = True
    r_ban.font.color.rgb = RGBColor(15, 23, 42)

    p_theme_lbl = doc.add_paragraph()
    p_theme_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme_lbl.paragraph_format.space_before = Pt(5)
    p_theme_lbl.paragraph_format.space_after = Pt(4)
    r_tl = p_theme_lbl.add_run("THÈME :")
    r_tl.font.name = "Times New Roman"
    r_tl.font.size = Pt(13)
    r_tl.bold = True
    r_tl.font.color.rgb = RGBColor(180, 0, 0)

    # Title box (table single cell with elegant border & fill)
    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_tbl.cell(0, 0)
    cell.width = Cm(16.0)
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    set_cell_shading(cell, "F8FAFC")
    set_cell_borders(cell, top="1E3A8A", bottom="1E3A8A", left="1E3A8A", right="1E3A8A", sz="12")
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.3
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    r_tit = p_title.add_run("InsurFlow : Conception et Réalisation d'une Plateforme ERP Cloud Native & Intelligente pour la Gestion Intégrale de la Production d'Assurance, du Double Circuit de Règlements, de la Facturation et des Exercices Comptables")
    r_tit.font.name = "Times New Roman"
    r_tit.font.size = Pt(14)
    r_tit.bold = True
    r_tit.font.color.rgb = RGBColor(30, 58, 138)

    p_sp1 = doc.add_paragraph()
    p_sp1.paragraph_format.space_before = Pt(20)

    # Student & Tutors Table
    info_tbl = doc.add_table(rows=3, cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_tbl.autofit = False
    for row in info_tbl.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
        for c in row.cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            set_cell_borders(c, top="FFFFFF", bottom="FFFFFF", left="FFFFFF", right="FFFFFF", sz="0")

    c00 = info_tbl.cell(0, 0).paragraphs[0]
    c00.paragraph_format.space_after = Pt(2)
    r = c00.add_run("Réalisé par :\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(180, 0, 0)
    r2 = c00.add_run("Élève-Ingénieur en 4IIR")
    r2.font.size = Pt(11)
    r2.bold = True

    c01 = info_tbl.cell(0, 1).paragraphs[0]
    c01.paragraph_format.space_after = Pt(2)
    r = c01.add_run("Organisme d'accueil :\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(180, 0, 0)
    r2 = c01.add_run("Cabinet de Courtage en Assurances\nYK Software Solutions")
    r2.font.size = Pt(11)

    c10 = info_tbl.cell(1, 0).paragraphs[0]
    c10.paragraph_format.space_before = Pt(8)
    c10.paragraph_format.space_after = Pt(2)
    r = c10.add_run("Encadrant Pédagogique :\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(180, 0, 0)
    r2 = c10.add_run("Pr. [Nom & Prénom]\nEnseignant-Chercheur EMSI")
    r2.font.size = Pt(11)

    c11 = info_tbl.cell(1, 1).paragraphs[0]
    c11.paragraph_format.space_before = Pt(8)
    c11.paragraph_format.space_after = Pt(2)
    r = c11.add_run("Encadrant Professionnel :\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(180, 0, 0)
    r2 = c11.add_run("M. [Nom & Prénom]\nExpert Métier Assurance")
    r2.font.size = Pt(11)

    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bot.paragraph_format.space_before = Pt(25)
    p_bot.paragraph_format.space_after = Pt(0)
    r_an = p_bot.add_run("Année Universitaire : 2025 / 2026")
    r_an.font.name = "Times New Roman"
    r_an.font.size = Pt(12)
    r_an.bold = True
    r_an.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_page_break()

    # 3. Parse Markdown body line by line
    lines = content.split('\n')
    i = 0
    total_lines = len(lines)

    # Skip until after the title metadata block
    while i < total_lines and not lines[i].startswith('# Dédicaces'):
        i += 1

    in_code_block = False
    code_lang = ""
    code_lines = []

    in_table = False
    table_lines = []

    while i < total_lines:
        line = lines[i]

        # Handle page breaks
        if line.strip() == '\\newpage':
            doc.add_page_break()
            i += 1
            continue

        # Handle code blocks (including Mermaid / PlantUML / Code / Diagrams)
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                # Render code block / diagram box
                if code_lines:
                    tbl = doc.add_table(rows=1, cols=1)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    c = tbl.cell(0, 0)
                    c.width = Cm(16.0)
                    set_cell_margins(c, top=120, bottom=120, left=140, right=140)
                    set_cell_shading(c, "F1F5F9")
                    set_cell_borders(c, top="CBD5E1", bottom="CBD5E1", left="3B82F6", right="CBD5E1", sz="8")
                    
                    cp = c.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cp.paragraph_format.line_spacing = 1.0
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(2)
                    cp.paragraph_format.first_line_indent = Cm(0)
                    
                    code_text = '\n'.join(code_lines)
                    c_run = cp.add_run(code_text)
                    c_run.font.name = "Consolas"
                    c_run.font.size = Pt(8.0 if len(code_lines) > 25 else 8.5)
                    c_run.font.color.rgb = RGBColor(30, 41, 59)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle Markdown Tables
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                # Process and render table
                if len(table_lines) >= 2:
                    raw_rows = []
                    for t_line in table_lines:
                        if re.match(r'^\s*\|(?:\s*:?-+:?\s*\|)+\s*$', t_line):
                            continue
                        cells_data = [c.strip() for c in t_line.strip().split('|')[1:-1]]
                        raw_rows.append(cells_data)

                    if raw_rows:
                        num_cols = max(len(r) for r in raw_rows)
                        num_rows = len(raw_rows)
                        tbl = doc.add_table(rows=num_rows, cols=num_cols)
                        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                        tbl.autofit = False

                        # Column width presets based on column count
                        col_widths = []
                        font_size_table = 9.5
                        pad_lr = 100
                        pad_tb = 80

                        is_planning_10 = (num_cols == 10 and any(
                            any(k in str(c).lower() for k in ['juil', 'août', 'semaine', 'phase', 'étape', 'planning'])
                            for c in raw_rows[0]
                        ))
                        is_planning_9 = (num_cols == 9 and any(
                            any(k in str(c).lower() for k in ['juil', 'août', 'semaine', 'phase', 'étape', 'planning'])
                            for c in raw_rows[0]
                        ))
                        is_planning = is_planning_10 or is_planning_9
                        is_data_dict = (num_cols == 9 and not is_planning)

                        if num_cols == 2:
                            col_widths = [Cm(3.8), Cm(12.2)]
                            font_size_table = 10.0
                        elif num_cols == 3:
                            col_widths = [Cm(4.0), Cm(6.0), Cm(6.0)]
                            font_size_table = 9.5
                        elif is_planning_10: # Planning table (Phase + 9 weeks: Juil S1..S4, Août S1..S5)
                            col_widths = [Cm(6.55)] + [Cm(1.05)] * 9
                            font_size_table = 8.5
                            pad_lr = 30
                            pad_tb = 50
                        elif is_planning_9: # Planning table (Phase + 8 weeks: Juil S1..S4, Août S1..S4)
                            col_widths = [Cm(8.0)] + [Cm(1.0)] * 8
                            font_size_table = 9.0
                            pad_lr = 30
                            pad_tb = 55
                        elif is_data_dict: # Data dictionary
                            col_widths = [Cm(2.2), Cm(1.9), Cm(0.9), Cm(1.2), Cm(1.5), Cm(2.8), Cm(1.1), Cm(1.1), Cm(2.3)]
                            font_size_table = 7.5
                            pad_lr = 40
                            pad_tb = 50
                        elif num_cols == 17: # Old planning table fallback
                            col_widths = [Cm(4.8)] + [Cm(0.7)] * 16
                            font_size_table = 7.5
                            pad_lr = 40
                            pad_tb = 50
                        elif num_cols >= 7:
                            w = 16.0 / num_cols
                            col_widths = [Cm(w)] * num_cols
                            font_size_table = 8.0
                            pad_lr = 50
                            pad_tb = 60
                        else:
                            w = 16.0 / num_cols
                            col_widths = [Cm(w)] * num_cols

                        for r_idx, r_data in enumerate(raw_rows):
                            row = tbl.rows[r_idx]
                            set_row_cant_split(row)
                            is_header = (r_idx == 0)
                            if is_header:
                                set_row_header(row)

                            for c_idx in range(num_cols):
                                cell_val = r_data[c_idx] if c_idx < len(r_data) else ""
                                cell = row.cells[c_idx]
                                if c_idx < len(col_widths):
                                    cell.width = col_widths[c_idx]
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                set_cell_margins(cell, top=pad_tb, bottom=pad_tb, left=pad_lr, right=pad_lr)

                                if is_header:
                                    set_cell_shading(cell, "1E3A8A") # EMSI / Tech Navy
                                    set_cell_borders(cell, top="1E3A8A", bottom="1E3A8A", left="CBD5E1", right="CBD5E1", sz="6")
                                else:
                                    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                                    set_cell_shading(cell, bg)
                                    set_cell_borders(cell, top="E2E8F0", bottom="E2E8F0", left="E2E8F0", right="E2E8F0", sz="4")

                                p = cell.paragraphs[0]
                                p.paragraph_format.line_spacing = 1.15
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                p.paragraph_format.first_line_indent = Cm(0)
                                
                                # Alignment in cell
                                if is_header:
                                    if is_planning and c_idx > 0:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    if re.match(r'^[0-9\s,\.DH%]+$', cell_val) or cell_val in ['X', 'Oui', 'Non']:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                                font_col = RGBColor(255, 255, 255) if is_header else (
                                    RGBColor(180, 0, 0) if (is_planning and cell_val == 'X') else RGBColor(15, 23, 42)
                                )
                                default_bold_val = is_header or (is_planning and cell_val == 'X')
                                f_size = font_size_table
                                if is_planning and is_header:
                                    f_size = 8.0
                                elif is_planning and cell_val == 'X':
                                    f_size = 10.0

                                add_formatted_text(p, cell_val, default_bold=default_bold_val, font_size=f_size, default_color=font_col)

                        # Space after table
                        p_sp = doc.add_paragraph()
                        p_sp.paragraph_format.space_before = Pt(2)
                        p_sp.paragraph_format.space_after = Pt(4)
                table_lines = []

        # Handle Figure / Table captions in HTML div or bold text
        if '<div align="center">' in line or '<div align="right">' in line:
            align = WD_ALIGN_PARAGRAPH.RIGHT if 'right' in line else WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            caption_text = ""
            while i < total_lines and '</div>' not in lines[i]:
                caption_text += lines[i] + " "
                i += 1
            i += 1 # skip </div>
            caption_text = re.sub(r'<[^>]+>', '', caption_text).strip()
            if caption_text:
                cp = doc.add_paragraph()
                cp.alignment = align
                cp.paragraph_format.line_spacing = 1.15
                cp.paragraph_format.first_line_indent = Cm(0)
                is_fig = caption_text.startswith('Figure')
                is_tab = caption_text.startswith('Tableau')
                
                # If it's a table title, it goes above the table (EMSI rule) -> space before 8, space after 4
                # If it's a figure caption, it goes below the figure (EMSI rule) -> space before 4, space after 10
                if is_tab:
                    cp.paragraph_format.space_before = Pt(10)
                    cp.paragraph_format.space_after = Pt(4)
                else:
                    cp.paragraph_format.space_before = Pt(4)
                    cp.paragraph_format.space_after = Pt(10)

                r = cp.add_run(caption_text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = is_fig or is_tab
                r.italic = not (is_fig or is_tab)
                if is_fig or is_tab:
                    r.font.color.rgb = RGBColor(30, 58, 138)
            continue

        # Handle Horizontal Rules ---
        if re.match(r'^\s*---\s*$', line):
            i += 1
            continue

        # Handle Headings
        if line.startswith('# '):
            h_text = line[2:].strip()
            clean_title = re.sub(r'<[^>]+>', '', h_text).strip()
            
            # Start chapters and major sections on new page
            if clean_title.startswith('Chapitre') or clean_title.startswith('Conclusion') or clean_title.startswith('Bibliographie') or clean_title.startswith('Annexes') or clean_title.startswith('Introduction'):
                if doc.paragraphs and len(doc.paragraphs) > 5:
                    doc.add_page_break()

            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_before = Pt(18)
            hp.paragraph_format.space_after = Pt(8)
            hp.paragraph_format.first_line_indent = Cm(0)
            hp.paragraph_format.keep_with_next = True
            
            r = hp.add_run(clean_title)
            r.font.name = "Times New Roman"
            r.font.size = Pt(16)
            r.bold = True
            r.font.color.rgb = RGBColor(180, 0, 0) if clean_title.startswith('Chapitre') else RGBColor(15, 23, 42)
            i += 1
            continue

        elif line.startswith('## '):
            h_text = line[3:].strip()
            clean_title = re.sub(r'<[^>]+>', '', h_text).strip()
            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_before = Pt(14)
            hp.paragraph_format.space_after = Pt(6)
            hp.paragraph_format.first_line_indent = Cm(0)
            hp.paragraph_format.keep_with_next = True
            
            r = hp.add_run(clean_title)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)
            i += 1
            continue

        elif line.startswith('### '):
            h_text = line[4:].strip()
            clean_title = re.sub(r'<[^>]+>', '', h_text).strip()
            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_before = Pt(10)
            hp.paragraph_format.space_after = Pt(4)
            hp.paragraph_format.first_line_indent = Cm(0)
            hp.paragraph_format.keep_with_next = True
            
            r = hp.add_run(clean_title)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(51, 65, 85)
            i += 1
            continue

        elif line.startswith('#### '):
            h_text = line[5:].strip()
            clean_title = re.sub(r'<[^>]+>', '', h_text).strip()
            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_before = Pt(8)
            hp.paragraph_format.space_after = Pt(3)
            hp.paragraph_format.first_line_indent = Cm(0)
            hp.paragraph_format.keep_with_next = True
            
            r = hp.add_run(clean_title)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
            r.font.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)
            i += 1
            continue

        # Handle Bullet lists
        if re.match(r'^\s*[-*]\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            bullet_text = re.sub(r'^\s*[-*]\s+', '', line).strip()
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            lp.paragraph_format.line_spacing = 1.3
            lp.paragraph_format.space_before = Pt(2)
            lp.paragraph_format.space_after = Pt(3)
            lp.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.5)
            lp.paragraph_format.first_line_indent = Cm(-0.4)
            
            r_b = lp.add_run("•  " if indent_level == 0 else "–  ")
            r_b.font.name = "Times New Roman"
            r_b.font.size = Pt(12)
            r_b.font.color.rgb = RGBColor(180, 0, 0) if indent_level == 0 else RGBColor(100, 116, 139)
            
            add_formatted_text(lp, bullet_text, font_size=12)
            i += 1
            continue

        # Handle Numbered lists (1. 2. 3.)
        if re.match(r'^\s*\d+\.\s+', line):
            num_match = re.match(r'^\s*(\d+\.)\s+(.*)', line)
            num_prefix = num_match.group(1)
            num_text = num_match.group(2)
            
            np_p = doc.add_paragraph()
            np_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            np_p.paragraph_format.line_spacing = 1.3
            np_p.paragraph_format.space_before = Pt(2)
            np_p.paragraph_format.space_after = Pt(3)
            np_p.paragraph_format.left_indent = Cm(1.0)
            np_p.paragraph_format.first_line_indent = Cm(-0.5)
            
            r_num = np_p.add_run(num_prefix + "  ")
            r_num.font.name = "Times New Roman"
            r_num.font.size = Pt(12)
            r_num.bold = True
            r_num.font.color.rgb = RGBColor(30, 41, 59)
            
            add_formatted_text(np_p, num_text, font_size=12)
            i += 1
            continue

        # Handle Math formulas block ($$...$$)
        if line.strip().startswith('$$') and line.strip().endswith('$$'):
            formula = line.strip()[2:-2].strip()
            mp = doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mp.paragraph_format.space_before = Pt(6)
            mp.paragraph_format.space_after = Pt(6)
            mp.paragraph_format.first_line_indent = Cm(0)
            r = mp.add_run(formula)
            r.font.name = "Cambria Math"
            r.font.size = Pt(12)
            r.italic = True
            r.bold = True
            r.font.color.rgb = RGBColor(30, 58, 138)
            i += 1
            continue

        # Handle Regular Paragraphs
        stripped = line.strip()
        if stripped:
            # Check for Pedagogical Note callout box
            if stripped.startswith('**Note Explicative Pédagogique :**') or stripped.startswith('Note Explicative Pédagogique :'):
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                c = tbl.cell(0, 0)
                c.width = Cm(16.0)
                set_cell_margins(c, top=140, bottom=140, left=180, right=140)
                set_cell_shading(c, "F8FAFC")
                set_cell_borders(c, top="E2E8F0", bottom="E2E8F0", left="1E3A8A", right="E2E8F0", sz="12")
                
                cp = c.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                cp.paragraph_format.line_spacing = 1.25
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                cp.paragraph_format.first_line_indent = Cm(0)
                
                add_formatted_text(cp, stripped, font_size=11, default_color=RGBColor(30, 41, 59))
                
                p_sp = doc.add_paragraph()
                p_sp.paragraph_format.space_before = Pt(2)
                p_sp.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            # Check for Screenshot Placeholder Box
            if 'Insérer ici la capture d\'écran' in stripped or 'Insérer ici l\'image' in stripped:
                clean_ph = re.sub(r'[`\[\]]', '', stripped).strip()
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                c = tbl.cell(0, 0)
                c.width = Cm(16.0)
                set_cell_margins(c, top=160, bottom=160, left=160, right=160)
                set_cell_shading(c, "EFF6FF") # Soft Blue Background
                set_cell_borders(c, top="3B82F6", bottom="3B82F6", left="3B82F6", right="3B82F6", sz="8")
                
                cp = c.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.line_spacing = 1.15
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                cp.paragraph_format.first_line_indent = Cm(0)
                
                r = cp.add_run(f"📷  [ {clean_ph} ]")
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = True
                r.italic = True
                r.font.color.rgb = RGBColor(30, 58, 138)
                
                p_sp = doc.add_paragraph()
                p_sp.paragraph_format.space_before = Pt(2)
                p_sp.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            is_dedicace = stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**')
            clean_p_text = stripped[1:-1] if is_dedicace else stripped
            
            p = doc.add_paragraph()
            style_paragraph(p, font_name="Times New Roman", font_size=12,
                            align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                            space_before=6, space_after=6,
                            first_line_indent=0.0 if is_dedicace else 0.8)
            
            add_formatted_text(p, clean_p_text, default_italic=is_dedicace, font_size=12)

        i += 1

    # Save document
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == "__main__":
    md_file = os.path.abspath(r"c:\Users\lenovo\Documents\GitHub\InsurFlow\docs\rapport_pfa_insurflow.md")
    docx_file = os.path.abspath(r"c:\Users\lenovo\Documents\GitHub\InsurFlow\docs\Rapport_PFA_InsurFlow_EMSI.docx")
    build_docx_report(md_file, docx_file)
